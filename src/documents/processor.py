import os
import sys
from pathlib import Path
from typing import List, Dict, Optional
import logging
import re

# Add project root to path
sys.path.append(str(Path(__file__).parent.parent.parent))

from src.documents.russian_nlp import RussianNLP

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.russian_nlp = RussianNLP()
    
    def process_document(self, file_path: str, metadata: Optional[Dict] = None) -> Dict:
        """Process document and extract text with metadata"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            print(f"📄 Processing file: {file_path.name}")
            
            # Extract text based on file type
            text_content = self._extract_text(file_path)
            
            if not text_content or len(text_content.strip()) < 10:
                print(f"⚠️ No meaningful content extracted from {file_path.name}")
                return {
                    'file_path': str(file_path),
                    'file_name': file_path.name,
                    'raw_text': "",
                    'normalized_text': "",
                    'chunks': [],
                    'chunk_count': 0,
                    'metadata': metadata or {}
                }
            
            print(f"✅ Extracted {len(text_content)} characters")
            
            # Умная очистка текста
            clean_text = self._smart_clean_text(text_content)
            print(f"🧹 Smart-cleaned text: {len(clean_text)} characters")
            print(f"📝 Clean text preview: {clean_text[:300]}...")
            
            # Process with Russian NLP
            normalized_text = self.russian_nlp.normalize_text(clean_text)
            chunks = self.russian_nlp.chunk_text(normalized_text, chunk_size=600, overlap=100)
            
            result = {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'raw_text': text_content,
                'normalized_text': normalized_text,
                'chunks': chunks,
                'chunk_count': len(chunks),
                'metadata': metadata or {}
            }
            
            print(f"✅ Created {len(chunks)} chunks")
            logger.info(f"✓ Processed {file_path.name}: {len(chunks)} chunks")
            return result
            
        except Exception as e:
            print(f"❌ Failed to process {file_path}: {e}")
            logger.error(f"Failed to process {file_path}: {e}")
            raise
    
    def _smart_clean_text(self, text: str) -> str:
        """Умная очистка текста сохраняя структуру"""
        if not text:
            return ""
        
        # Разбиваем на строки
        lines = text.split('\n')
        clean_lines = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Пропускаем очевидный мусор но сохраняем содержательные таблицы
            if self._is_junk_line(line):
                continue
            
            # Улучшаем форматирование таблиц
            if '|' in line and len(line) > 10:
                # Это строка таблицы - улучшаем форматирование
                parts = [p.strip() for p in line.split('|') if p.strip()]
                if len(parts) >= 2:  # Минимум 2 колонки
                    line = ' - '.join(parts)
            
            clean_lines.append(line)
        
        # Объединяем с правильными отступами
        result = '\n'.join(clean_lines)
        
        # Финальная очистка
        result = re.sub(r'\n{3,}', '\n\n', result)
        result = re.sub(r'\s+', ' ', result)
        result = result.strip()
        
        return result
    
    def _is_junk_line(self, line: str) -> bool:
        """Определяет мусорные строки но сохраняет полезное содержимое"""
        line_lower = line.lower().strip()
        
        # Пустые или очень короткие
        if len(line_lower) < 2:
            return True
        
        # Только цифры или символы
        if re.match(r'^[\d\.\-\s\|]+$', line_lower):
            return True
            
        # Строки состоящие только из разделителей
        if re.match(r'^[-=_\|\s]+$', line):
            return True
            
        return False
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from various file formats"""
        suffix = file_path.suffix.lower()
        
        print(f"🔍 Extracting text from {suffix} file...")
        
        if suffix == '.txt':
            return self._read_text_file(file_path)
        elif suffix == '.docx':
            return self._read_docx_file(file_path)
        elif suffix == '.pdf':
            return self._read_pdf_file(file_path)
        else:
            print(f"⚠️ Unsupported file type {suffix}, trying as text...")
            return self._read_text_file(file_path)
    
    def _read_text_file(self, file_path: Path) -> str:
        """Read plain text file with encoding detection"""
        encodings = ['utf-8', 'cp1251', 'koi8-r', 'iso-8859-5']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    print(f"✅ Successfully read with {encoding} encoding")
                    return content
            except UnicodeDecodeError:
                continue
        
        print(f"❌ Could not decode text file: {file_path}")
        return ""
    
    def _read_docx_file(self, file_path: Path) -> str:
        """Read DOCX file с улучшенной обработкой таблиц"""
        try:
            from docx import Document
            doc = Document(file_path)
            all_content = []
            
            # Читаем параграфы И таблицы в правильном порядке
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Параграф
                    # Находим соответствующий параграф
                    for para in doc.paragraphs:
                        if para._element == element:
                            text = para.text.strip()
                            if text and len(text) > 2:
                                all_content.append(text)
                            break
                
                elif element.tag.endswith('tbl'):  # Таблица
                    # Находим соответствующую таблицу
                    for table in doc.tables:
                        if table._element == element:
                            table_text = self._extract_table_text(table)
                            if table_text:
                                all_content.append(table_text)
                            break
            
            result = '\n\n'.join(all_content)
            print(f"✅ Extracted structured text from DOCX: {len(result)} characters")
            return result
            
        except ImportError:
            print("❌ python-docx not installed")
            return ""
        except Exception as e:
            print(f"❌ Failed to read DOCX: {e}")
            # Fallback to simple method
            return self._read_docx_simple(file_path)
    
    def _extract_table_text(self, table) -> str:
        """Извлекаем текст из таблицы в читаемом формате"""
        table_lines = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)
            
            if cells:
                # Для таблиц-инструкций делаем удобный формат
                if len(cells) == 2:  # Проблема -> Решение
                    table_lines.append(f"ПРОБЛЕМА: {cells[0]}")
                    table_lines.append(f"РЕШЕНИЕ: {cells[1]}")
                    table_lines.append("")  # Пустая строка для разделения
                else:
                    # Обычная таблица
                    table_lines.append(" | ".join(cells))
        
        return '\n'.join(table_lines)
    
    def _read_docx_simple(self, file_path: Path) -> str:
        """Простой метод чтения DOCX как fallback"""
        try:
            from docx import Document
            doc = Document(file_path)
            content_parts = []
            
            # Читаем параграфы
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and len(text) > 2:
                    content_parts.append(text)
            
            # Читаем таблицы простым способом
            for table in doc.tables:
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        content_parts.append(' | '.join(row_cells))
            
            result = '\n\n'.join(content_parts)
            print(f"✅ Extracted simple DOCX text: {len(result)} characters")
            return result
            
        except Exception as e:
            print(f"❌ Simple DOCX read also failed: {e}")
            return ""
    
    def _read_pdf_file(self, file_path: Path) -> str:
        """Read PDF file"""
        try:
            import PyPDF2
            
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text.strip())
            
            content = '\n\n'.join(text_parts)
            print(f"✅ Extracted text from PDF: {len(content)} characters")
            return content
            
        except ImportError:
            print("❌ PyPDF2 not installed")
            return ""
        except Exception as e:
            print(f"❌ Failed to read PDF: {e}")
            return ""
